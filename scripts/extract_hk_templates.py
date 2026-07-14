#!/usr/bin/env python3
"""从 Appendix D1–D9 模板提取字段清单（非全文入库）。"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document as DocxDocument
from openpyxl import load_workbook

from rag.industry_hk.extract_utils import sha256_file, utc_now, write_json
from rag.industry_hk.paths import (
    HK_SOURCES_MANIFEST_PATH,
    HK_TEMPLATES_DIR,
    TEMPLATE_SOURCES,
)

FIELD_RE = re.compile(r"^[A-Z][A-Za-z0-9 /\-()]{2,80}$")
SKIP_HEADERS = {
    "field",
    "fields",
    "item",
    "no",
    "number",
    "description",
    "name",
    "title",
    "section",
    "page",
}


def _normalize_field(text: str) -> str:
    text = re.sub(r"\s+", " ", text.strip())
    return text


def _extract_docx_fields(path: Path) -> list[str]:
    doc = DocxDocument(str(path))
    fields: list[str] = []
    seen: set[str] = set()

    def add(value: str) -> None:
        value = _normalize_field(value)
        if not value or len(value) < 3:
            return
        key = value.lower()
        if key in SKIP_HEADERS or key in seen:
            return
        if value.isdigit():
            return
        seen.add(key)
        fields.append(value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = _normalize_field(cell.text)
                if FIELD_RE.match(text):
                    add(text)
                for line in text.split("\n"):
                    line = _normalize_field(line)
                    if FIELD_RE.match(line):
                        add(line)

    for paragraph in doc.paragraphs:
        text = _normalize_field(paragraph.text)
        if not text:
            continue
        if text.endswith(":") and len(text) < 60:
            add(text[:-1])
        if FIELD_RE.match(text) and len(text.split()) <= 8:
            add(text)

    return fields[:120]


def _extract_xlsx_fields(path: Path) -> list[str]:
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    fields: list[str] = []
    seen: set[str] = set()
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(max_row=40, values_only=True):
            for value in row:
                if value is None:
                    continue
                text = _normalize_field(str(value))
                key = text.lower()
                if not text or key in seen or key in SKIP_HEADERS:
                    continue
                if FIELD_RE.match(text):
                    seen.add(key)
                    fields.append(text)
    return fields[:120]


def _render_template_md(spec: dict, fields: list[str]) -> str:
    num = spec["template_num"]
    name = spec["name"]
    lines = [
        "---",
        f"source_file: {spec['path'].relative_to(PROJECT_ROOT)}",
        f"doc_id: template_d{num}",
        f"section_id: d{num}_fields",
        f"title: Appendix D{num} Field Reference ({name})",
        f"authority: CICBIMS 2024 Appendix D{num}",
        "priority: high",
        "language: en",
        "---",
        "",
        f"# Appendix D{num} — {name} (Field Reference)",
        "",
        "Extracted field / table headers only. Full template not ingested.",
        "",
        "## Fields",
        "",
    ]
    for field in fields:
        lines.append(f"- {field}")
    lines.append("")
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    HK_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    template_manifest_rows: list[dict] = []

    for spec in TEMPLATE_SOURCES:
        path = Path(spec["path"])
        if not path.is_file():
            print(f"缺少模板：{path}", file=sys.stderr)
            return 1

        if path.suffix.lower() == ".docx":
            fields = _extract_docx_fields(path)
        elif path.suffix.lower() == ".xlsx":
            fields = _extract_xlsx_fields(path)
        else:
            fields = []

        if not fields:
            fields = [f"D{spec['template_num']} template placeholder field"]

        out_path = HK_TEMPLATES_DIR / f"D{spec['template_num']}_fields.md"
        out_path.write_text(_render_template_md(spec, fields), encoding="utf-8")
        stat = path.stat()
        template_manifest_rows.append(
            {
                "doc_id": spec["doc_id"],
                "path": str(path.relative_to(PROJECT_ROOT)),
                "sha256": sha256_file(path),
                "byte_size": stat.st_size,
                "sheet_or_para_count": len(fields),
                "mtime": stat.st_mtime,
                "kind": "template",
                "output_md": str(out_path.relative_to(PROJECT_ROOT)),
            }
        )
        print(f"D{spec['template_num']}: {len(fields)} fields -> {out_path.name}")

    if HK_SOURCES_MANIFEST_PATH.is_file():
        manifest = json.loads(HK_SOURCES_MANIFEST_PATH.read_text(encoding="utf-8"))
    else:
        manifest = {"sources": []}
    pdf_sources = [
        item for item in manifest.get("sources", []) if item.get("kind") == "pdf"
    ]
    manifest["sources"] = pdf_sources + template_manifest_rows
    manifest["templates_updated_at"] = utc_now()
    write_json(HK_SOURCES_MANIFEST_PATH, manifest)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
